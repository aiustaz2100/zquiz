import os
import json
import random
import string
import io
import base64
import requests
import re
import urllib.request
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
from openai import OpenAI
from dotenv import load_dotenv

# --- ЖАҢА КІТАПХАНАЛАР ---
try:
    import qrcode
    from docx import Document
    from PyPDF2 import PdfReader
    from youtube_transcript_api import YouTubeTranscriptApi
    from bs4 import BeautifulSoup
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    print("⚠️ Кейбір кітапханалар жоқ.")

basedir = os.path.abspath(os.path.dirname(__file__))
env_path = os.path.join(basedir, '.env')
if os.path.exists(env_path):
    load_dotenv(env_path)

app = Flask(__name__)
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

TEST_STORAGE = {}
MAX_FILE_SIZE = 2 * 1024 * 1024  

# --- КӨМЕКШІ ФУНКЦИЯЛАР ---

def extract_text_from_file(file):
    filename = file.filename.lower()
    content = ""
    
    if filename.endswith('.pdf'):
        reader = PdfReader(file)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                content += extracted + "\n"
        
        if not content.strip():
            raise ValueError("Бұл PDF файлдан мәтін табылмады (мүмкін ол сканерленген сурет). Басқа мәтіні бар құжат жүктеп көріңіз.")
            
    elif filename.endswith('.docx'):
        doc = Document(file)
        for para in doc.paragraphs:
            content += para.text + "\n"
            
    elif filename.endswith('.txt'):
        content = file.read().decode('utf-8')

    return content[:15000]

def get_youtube_transcript(url):
    try:
        # Улучшенный поиск ID видео (понимает любые ссылки, включая Shorts)
        video_id_match = re.search(r"(?:v=|\/|youtu\.be\/|embed\/)([0-9A-Za-z_-]{11})", url)
        if not video_id_match:
            return None, "Сілтеме қате немесе YouTube видеосы емес."
        
        video_id = video_id_match.group(1)
        
        try:
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        except Exception:
            return None, "Бұл видеода ешқандай субтитр жоқ (тіпті автоматты түрде де)."

        # Сначала ищем нужные языки
        try:
            transcript = transcript_list.find_transcript(['kk', 'ru', 'en']).fetch()
        except:
            # Если их нет, берем вообще ПЕРВЫЙ попавшийся язык (План Б)
            try:
                transcript = list(transcript_list)[0].fetch()
            except:
                return None, "Видеодан мәтін алу мүмкін болмады."

        text = " ".join([t['text'] for t in transcript])
        return text[:15000], None
    except Exception as e:
        print(f"YouTube Error: {e}")
        return None, f"Қате шықты: {str(e)}"

def get_url_content(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        resp = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(resp.content, 'html.parser')
        
        paragraphs = soup.find_all('p')
        text = " ".join([p.get_text() for p in paragraphs])
        return text[:15000]
    except:
        return None

# --- ROUTES ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_quiz', methods=['POST'])
def generate_quiz():
    try:
        if not api_key: return jsonify({"error": "API Key табылмады"}), 500

        topic = request.form.get('topic', '')
        grade = request.form.get('grade', 'Any')
        language = request.form.get('language', 'Kazakh')
        quiz_type = request.form.get('type', 'Multiple Choice')
        input_type = request.form.get('free_input_type', 'text')
        mode = request.form.get('mode', 'live')  
        
        context_text = ""
        image_data = None 

        if input_type == 'file' and 'file_upload' in request.files:
            file = request.files['file_upload']
            if file.filename != '':
                file.seek(0, os.SEEK_END)
                size = file.tell()
                file.seek(0)
                
                if size > MAX_FILE_SIZE:
                    return jsonify({"error": "Файл 2 МБ-тан аспауы керек!"}), 400

                if file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                    image_data = base64.b64encode(file.read()).decode('utf-8')
                    context_text = "Analyze this image and create a quiz based on it."
                else:
                    try:
                        context_text = extract_text_from_file(file)
                    except Exception as e:
                        return jsonify({"error": str(e)}), 400

        elif input_type == 'url':
            url_link = request.form.get('url_link')
            if url_link:
                extracted = get_url_content(url_link)
                if extracted:
                    context_text = f"Content from URL: {extracted}"
                else:
                    return jsonify({"error": "Сайттан мәтін оқылмады"}), 400

        elif input_type == 'youtube':
            yt_link = request.form.get('youtube_link')
            if yt_link:
                # Используем новую логику с отловом ошибок
                transcript, yt_error = get_youtube_transcript(yt_link)
                if yt_error:
                    return jsonify({"error": yt_error}), 400
                if transcript:
                    context_text = f"Video Transcript: {transcript}"
        
        if not context_text:
            context_text = topic

        raw_count = request.form.get('count', '5')
        q_count = random.randint(5, 30) if raw_count == 'auto' else int(raw_count)
        
        timer_req = request.form.get('timer', 'no')
        timer_val = (q_count * 60) if timer_req == 'yes' else 0

        otp = ''.join(random.choices(string.digits, k=4))

        system_instruction = f"""
        You are a Quiz Generator. 
        Create {q_count} questions based on the provided content/topic.
        Target: {grade} grade. Language: {language}.
        Format: JSON.
        CRITICAL: You MUST provide exactly 4 options for EVERY question.
        Structure: {{ "title": "...", "questions": [ {{ "question": "...", "options": ["Option 1", "Option 2", "Option 3", "Option 4"], "answer": "Option 1" }} ] }}
        """

        messages = [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": context_text[:20000]} 
        ]

        if image_data:
            messages[1] = {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Create a quiz based on this image content."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
                ]
            }

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            response_format={"type": "json_object"}
        )

        quiz_json = json.loads(response.choices[0].message.content)

        room_status = "started" if mode == "homework" else "waiting"

        TEST_STORAGE[otp] = {
            "id": otp,
            "data": quiz_json,
            "config": {"timer": timer_val, "topic": topic[:50]},
            "status": room_status,
            "mode": mode,
            "players": [],
            "scores": {} 
        }

        return jsonify({"status": "success", "otp": otp})

    except Exception as e:
        print(f"ERR: {e}")
        return jsonify({"error": str(e)}), 500


@app.template_filter('qrcode')
def qrcode_filter(data):
    qr = qrcode.QRCode(box_size=10, border=5); qr.add_data(data); qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white"); buf = io.BytesIO(); img.save(buf)
    return base64.b64encode(buf.getvalue()).decode('utf-8')

@app.route('/room/<otp>')
def room(otp):
    test = TEST_STORAGE.get(otp)
    if not test: return "Not Found"
    mode = test.get('mode', 'live')
    return render_template('quiz_room.html', test=test, otp=otp, join_url=request.host_url+"join/"+otp, mode=mode)

@app.route('/get_players/<otp>')
def get_players(otp):
    test = TEST_STORAGE.get(otp)
    if not test: return jsonify({"error": "Not Found"}), 404
    
    if test['status'] != 'waiting':
        leaderboard = []
        for p in test['players']:
            score = test['scores'].get(p, 0) 
            leaderboard.append({"name": p, "score": score})
        return jsonify({"players": leaderboard, "status": test['status']})
    
    return jsonify({"players": test['players'], "status": test['status']})

@app.route('/start_quiz/<otp>')
def start_quiz(otp):
    if otp in TEST_STORAGE: TEST_STORAGE[otp]['status'] = 'started'; return jsonify({"status": "started"})
    return "Err", 404

@app.route('/join/<otp>', methods=['GET', 'POST'])
def student_join(otp):
    test = TEST_STORAGE.get(otp)
    if not test: return "Err"
    if request.method == 'POST':
        nick = request.form.get('nickname')
        if nick: 
            if nick not in test['players']: test['players'].append(nick)
            return redirect(url_for('take_test', otp=otp, nickname=nick))
    return render_template('student_login.html', otp=otp, topic=test['data']['title'])

@app.route('/test/<otp>/<nickname>')
def take_test(otp, nickname):
    test = TEST_STORAGE.get(otp)
    if not test: return "Err"
    if test['status'] == 'waiting': return render_template('waiting_room.html', otp=otp, nickname=nickname)
    return render_template('student_test.html', test=test, nickname=nickname, otp=otp, timer=test['config']['timer'])

@app.route('/submit_test/<otp>', methods=['POST'])
def submit_test(otp):
    test = TEST_STORAGE.get(otp)
    nick = request.form.get('nickname')
    questions = test['data']['questions']; score = 0
    
    for i, q in enumerate(questions):
        if request.form.get(f'q{i}') == q['answer']: score += 1
        
    test['scores'][nick] = score
        
    return render_template('student_result.html', score=score, total=len(questions), student_name=nick)

@app.route('/download/<otp>/<fmt>')
def download_file(otp, fmt):
    test = TEST_STORAGE.get(otp)
    if not test: return "Not Found", 404
    
    title = test['data'].get('title', 'ZQuiz')
    questions = test['data'].get('questions', [])
    
    if fmt == 'docx':
        doc = Document()
        doc.add_heading(title, 0)
        
        for i, q in enumerate(questions):
            doc.add_paragraph(f"{i+1}. {q.get('question', '')}", style='List Number')
            for opt in q.get('options', []):
                doc.add_paragraph(opt, style='List Bullet')
            doc.add_paragraph(f"Дұрыс жауап: {q.get('answer', '')}")
            doc.add_paragraph() 
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream, 
            as_attachment=True, 
            download_name=f"quiz_{otp}.docx", 
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    elif fmt == 'pdf':
        
        # Скачиваем шрифт во ВРЕМЕННУЮ папку (/tmp/), к которой у Render точно есть доступ
        font_path = "/tmp/Roboto-Regular.ttf"
        if not os.path.exists(font_path):
            try:
                urllib.request.urlretrieve("https://raw.githubusercontent.com/googlefonts/roboto/main/src/hinted/Roboto-Regular.ttf", font_path)
            except Exception as e:
                print("Font download error:", e)

        file_stream = io.BytesIO()
        styles = getSampleStyleSheet()
        
        try:
            pdfmetrics.registerFont(TTFont('Roboto', font_path))
            style_n = styles['Normal']
            style_h = styles['Heading1']
            style_n.fontName = 'Roboto'
            style_h.fontName = 'Roboto'
        except Exception as e:
            print("Could not register font:", e)
            style_n = styles['Normal']
            style_h = styles['Heading1']
        
        story = []
        story.append(Paragraph(title, style_h))
        story.append(Spacer(1, 12))
        
        for i, q in enumerate(questions):
            story.append(Paragraph(f"<b>{i+1}. {q.get('question', '')}</b>", style_n))
            for opt in q.get('options', []):
                story.append(Paragraph(f"- {opt}", style_n))
            story.append(Paragraph(f"<i>Дұрыс жауап: {q.get('answer', '')}</i>", style_n))
            story.append(Spacer(1, 12))
            
        doc = SimpleDocTemplate(file_stream, pagesize=letter)
        doc.build(story)
        file_stream.seek(0)
        
        return send_file(
            file_stream, 
            as_attachment=True, 
            download_name=f"quiz_{otp}.pdf", 
            mimetype="application/pdf"
        )
        
    return "Формат табылған жоқ", 400

if __name__ == '__main__':
    app.run(debug=True)
